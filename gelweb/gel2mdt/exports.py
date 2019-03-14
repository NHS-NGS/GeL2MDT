"""Copyright (c) 2018 Great Ormond Street Hospital for Children NHS Foundation
Trust & Birmingham Women's and Children's NHS Foundation Trust

Permission is hereby granted, free of charge, to any person obtaining a copy of
this software and associated documentation files (the "Software"), to deal in
the Software without restriction, including without limitation the rights to
use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies
of the Software, and to permit persons to whom the Software is furnished to do
so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""
from .models import *
import csv
import xlsxwriter
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
import os
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def write_mdt_export(mdt_instance, mdt_reports):
    '''
    Writes a summary of the cases which are being brought to MDT
    :param writer: CSV file writer
    :param mdt_instance:  MDT instance
    :param mdt_reports: List of reports which are present in MDT
    :return: CSV file Writer
    '''
    failed_reports = []
    for report in mdt_reports:
        proband_variants = ProbandVariant.objects.filter(interpretation_report=report.interpretation_report)
        for proband_variant in proband_variants:
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant is None:
                failed_reports.append(report.interpretation_report.ir_family.ir_family_id)

    if failed_reports:
        failed_reports_formatted = ' '.join(list(set(failed_reports)))
        raise ValueError(f"Transcripts have not been selected for the following reports: {failed_reports_formatted}")

    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet()
    # set formatting
    header_format = workbook.add_format({'bold': 1})
    vcenter_format = workbook.add_format({'valign': 'vcenter'})
    vcenter_date_format = workbook.add_format({'valign': 'vcenter', 'num_format': 'mm/dd/yyyy'})
    # write headings
    if mdt_instance.sample_type == 'raredisease':
        worksheet.write('A1', 'GEL ID', header_format)
        worksheet.write('B1', 'CIP ID', header_format)
        worksheet.write('C1', 'GMC', header_format)
        worksheet.write('D1', 'Forename', header_format)
        worksheet.write('E1', 'Surname', header_format)
        worksheet.write('F1', 'Sex', header_format)
        worksheet.write('G1', 'DOB', header_format)
        worksheet.write('H1', 'NHS number', header_format)
        worksheet.write('I1', 'Family ID', header_format)
        worksheet.write('J1', 'Clinician', header_format)
        worksheet.write('K1', 'Panel(s)', header_format)
        worksheet.write('L1', 'Variant', header_format)
        worksheet.write('M1', 'Inheritance', header_format)
        worksheet.write('N1', 'Proband zygosity', header_format)
        worksheet.write('O1', 'Maternal zygosity', header_format)
        worksheet.write('P1', 'Paternal zygosity', header_format)
        worksheet.write('Q1', 'Phenotypic fit', header_format)
        worksheet.write('R1', 'Discussion required', header_format)
        worksheet.write('S1', 'Comments', header_format)
        worksheet.set_column('A:A', 10)
        worksheet.set_column('B:B', 10)
        worksheet.set_column('C:C', 10)
        worksheet.set_column('D:F', 10)
        worksheet.set_column('G:G', 10)
        worksheet.set_column('H:H', 10)
        worksheet.set_column('I:I', 10)
        worksheet.set_column('J:J', 10)
        worksheet.set_column('K:L', 40)
        worksheet.set_column('M:M', 10)
        worksheet.set_column('N:N', 18)
        worksheet.set_column('O:O', 18)
        worksheet.set_column('P:P', 30)
        worksheet.set_column('Q:Q', 30)
        worksheet.set_column('R:R', 30)
        worksheet.set_column('S:S', 30)
    elif mdt_instance.sample_type == 'cancer':
        worksheet.write('A1', 'GEL ID', header_format)
        worksheet.write('B1', 'CIP ID', header_format)
        worksheet.write('C1', 'LDP', header_format)
        worksheet.write('D1', 'Forename', header_format)
        worksheet.write('E1', 'Surname', header_format)
        worksheet.write('F1', 'Sex', header_format)
        worksheet.write('G1', 'DOB', header_format)
        worksheet.write('H1', 'NHS number', header_format)
        worksheet.write('I1', 'Family ID', header_format)
        worksheet.write('J1', 'Clinician', header_format)
        worksheet.write('K1', 'Recruiting Disease', header_format)
        worksheet.write('L1', 'Disease subtype', header_format)
        worksheet.set_column('A:A', 10)
        worksheet.set_column('B:B', 10)
        worksheet.set_column('C:C', 10)
        worksheet.set_column('D:F', 10)
        worksheet.set_column('G:G', 10)
        worksheet.set_column('H:H', 10)
        worksheet.set_column('I:I', 10)
        worksheet.set_column('J:J', 10)
        worksheet.set_column('K:L', 10)

    row_count = 2
    if mdt_instance.sample_type == 'cancer':
        for report in mdt_reports:
            worksheet.write('A' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.gel_id,
                            vcenter_format)
            worksheet.write('B' + str(row_count),
                            report.interpretation_report.ir_family.ir_family_id, vcenter_format)
            worksheet.write('C' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.gmc, vcenter_format)
            worksheet.write('D' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.forename,
                            vcenter_format)
            worksheet.write('E' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.surname,
                            vcenter_format)
            worksheet.write('F' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.sex, vcenter_format)
            worksheet.write('G' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.date_of_birth.date(),
                            vcenter_date_format)
            worksheet.write('H' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.nhs_number,
                            vcenter_format)
            worksheet.write('I' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.gel_family_id, vcenter_format)
            worksheet.write('J' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.clinician.name,
                            vcenter_format)
            worksheet.write('K' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.recruiting_disease,
                            vcenter_format)
            worksheet.write('L' + str(row_count),
                            report.interpretation_report.ir_family.participant_family.proband.disease_subtype,
                            vcenter_format)
            row_count += 1
    elif mdt_instance.sample_type == 'raredisease':
        for report in mdt_reports:
            proband_variants = ProbandVariant.objects.filter(interpretation_report=report.interpretation_report)
            panels = InterpretationReportFamilyPanel.objects.filter(ir_family=report.interpretation_report.ir_family)
            pv_output = []
            for proband_variant in proband_variants:
                transcript = proband_variant.get_transcript()
                transcript_variant = proband_variant.get_transcript_variant()
                if transcript and transcript_variant:
                    hgvs_c = None
                    hgvs_p = None
                    hgvs_c_split = transcript_variant.hgvs_c.split(':')
                    hgvs_p_split = transcript_variant.hgvs_p.split(':')
                    if len(hgvs_c_split) > 1:
                        hgvs_c = hgvs_c_split[1]
                    if len(hgvs_p_split) > 1:
                        hgvs_p = hgvs_p_split[1].replace('%3D', '=')
                    pv_output.append({'variant': f'{transcript.gene}, '
                                     f'{hgvs_c}, '
                                     f'{hgvs_p}, ',
                                     'inheritance': f'{proband_variant.inheritance}',
                                     'proband_zygosity': f'{proband_variant.zygosity}',
                                     'mat_zygosity': f'{proband_variant.maternal_zygosity}',
                                     'pat_zygosity': f'{proband_variant.paternal_zygosity}',})
            panel_names = []
            for panel in panels:
                panel_names.append(f'{panel.panel.panel.panel_name}_'
                                   f'{panel.panel.version_number}')

            v_rows = row_count
            for variant in pv_output:
                worksheet.write('L' + str(v_rows), variant['variant'])
                worksheet.write('M' + str(v_rows), variant['inheritance'])
                worksheet.write('N' + str(v_rows), variant['proband_zygosity'])
                worksheet.write('O' + str(v_rows), variant['mat_zygosity'])
                worksheet.write('P' + str(v_rows), variant['pat_zygosity'])
                worksheet.data_validation('Q' + str(v_rows),
                    {'validate': 'list', 'source': ['Yes', 'No', 'Maybe']})
                worksheet.data_validation('R' + str(v_rows),
                    {'validate': 'list', 'source': ['Yes', 'No']})
                v_rows += 1

            if row_count == v_rows - 1:
                worksheet.write('A' + str(row_count),
                                report.interpretation_report.ir_family.participant_family.proband.gel_id, vcenter_format)
                worksheet.write('B' + str(row_count),
                                report.interpretation_report.ir_family.ir_family_id, vcenter_format)
                worksheet.write('C' + str(row_count),
                                report.interpretation_report.ir_family.participant_family.proband.gmc, vcenter_format)
                worksheet.write('D' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.proband.forename, vcenter_format)
                worksheet.write('E' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.proband.surname, vcenter_format)
                worksheet.write('F' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.proband.sex, vcenter_format)
                worksheet.write('G' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.proband.date_of_birth.date(), vcenter_date_format)
                worksheet.write('H' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.proband.nhs_number, vcenter_format)
                worksheet.write('I' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.gel_family_id, vcenter_format)
                worksheet.write('J' + str(row_count),
                    report.interpretation_report.ir_family.participant_family.clinician.name, vcenter_format)
                worksheet.write('K' + str(row_count),
                    '\n'.join(panel_names), vcenter_format)
            else:
                worksheet.merge_range('A' + str(row_count) + ':A' + str(v_rows - 1),
                                      report.interpretation_report.ir_family.participant_family.proband.gel_id,
                                      vcenter_format)
                worksheet.merge_range('B' + str(row_count) + ':B' + str(v_rows - 1),
                                      report.interpretation_report.ir_family.ir_family_id,
                                      vcenter_format)
                worksheet.merge_range('C' + str(row_count) + ':C' + str(v_rows - 1),
                                      report.interpretation_report.ir_family.participant_family.proband.gmc,
                                      vcenter_format)
                worksheet.merge_range('D' + str(row_count) + ':D' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.proband.forename,
                    vcenter_format)
                worksheet.merge_range('E' + str(row_count) + ':E' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.proband.surname,
                    vcenter_format)
                worksheet.merge_range('F' + str(row_count) + ':F' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.proband.sex,
                    vcenter_format)
                worksheet.merge_range('G' + str(row_count) + ':G' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.proband.date_of_birth.date(),
                    vcenter_date_format)
                worksheet.merge_range('H' + str(row_count) + ':H' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.proband.nhs_number,
                    vcenter_format)
                worksheet.merge_range('I' + str(row_count) + ':I' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.gel_family_id,
                    vcenter_format)
                worksheet.merge_range('J' + str(row_count) + ':J' + str(v_rows-1),
                    report.interpretation_report.ir_family.participant_family.clinician.name,
                    vcenter_format)
                worksheet.merge_range('K' + str(row_count) + ':K' + str(v_rows-1),
                    '\n'.join(panel_names),
                    vcenter_format)

            row_count = v_rows

    workbook.close()
    # rewind the buffer
    output.seek(0)
    return output

def write_mdt_outcome_template(report):
    """
    :param pk: GEL Interpretationreport instance
    :return: Writes a docx template file for summarising proband MDT outcomes
    """
    document = Document()
    document.add_picture(os.path.join(settings.STATIC_DIR, 'nhs_image.png'))
    header_image = document.paragraphs[-1]
    header_image.alignment = WD_ALIGN_PARAGRAPH.RIGHT 
 
    document.add_heading('Genomics MDM record', 0)

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'THIS IS NOT A DIAGNOSTIC REPORT. UNVALIDATED FINDINGS SHOULD NOT BE USED TO INFORM CLINICAL '
        'MANAGEMENT DECISIONS.\n')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'This is a record of unvalidated variants identified through the 100,000 genome project. '
        'Class 3 variants are of uncertain clinical significance, future review and diagnostic confirmation '
        'may be appropriate if further evidence becomes available.\n')
    run.font.color.rgb = RGBColor(255, 0, 0)

    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.3)
    # table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.1)
    paragraph = document.add_paragraph()
    paragraph.add_run()

    table = document.add_table(rows=2, cols=4, style='Table Grid')
    heading_cells = table.rows[0].cells
    heading_cells[0].paragraphs[0].add_run('Patient Name').bold=True
    heading_cells[1].paragraphs[0].add_run('DOB').bold=True
    heading_cells[2].paragraphs[0].add_run('NHS number').bold=True
    heading_cells[3].paragraphs[0].add_run('Local ID').bold=True

    row = table.rows[1].cells
    row[0].text = str(report.ir_family.participant_family.proband.forename) \
                  + ' ' \
                  + str(report.ir_family.participant_family.proband.surname)
    row[1].text = str(report.ir_family.participant_family.proband.date_of_birth.date())
    try:
        row[2].text = report.ir_family.participant_family.proband.nhs_number
    except TypeError:
        row[2].text = ''
    if report.ir_family.participant_family.proband.local_id:
        row[3].text = report.ir_family.participant_family.proband.local_id

    paragraph = document.add_paragraph()
    paragraph.add_run()
    paragraph.add_run('Referring Clinician: ').bold=True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.clinician))
    paragraph.add_run('Department/Hospital: ').bold=True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.proband.gmc))
    paragraph.add_run('Study: ').bold=True
    paragraph.add_run('100,000 genomes (whole genome sequencing)\n')
    paragraph.add_run('OPA ID: ').bold = True
    paragraph.add_run('{}\n'.format(report.ir_family.ir_family_id))
    paragraph.add_run('Family ID: ').bold=True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.gel_family_id))
    paragraph.add_run('Genome Build: ').bold=True
    paragraph.add_run('{}\n\n'.format(report.assembly))

    # paragraph.add_run('Phenotype summary: ').bold=True
    # if sample_info.hpo_terms:
    #    paragraph.add_run('{}\n'.format(', '.join(list(json.loads(sample_info.hpo_terms)))))

    proband_variants = list(ProbandVariant.objects.filter(interpretation_report=report))

    run = paragraph.add_run('MDT:\n')
    run.font.size = Pt(16)
    run.underline = True
    run.bold = True

    if proband_variants:
        run = paragraph.add_run('Variant Outcome Summary:\n')
        run.font.size = Pt(13)
        run.underline = True
        run.bold = True

        table = document.add_table(rows=1, cols=7, style='Table Grid')
        heading_cells = table.rows[0].cells
        run = heading_cells[0].paragraphs[0].add_run('Gene')
        run.bold=True
        run.font.size = Pt(9)
        run = heading_cells[1].paragraphs[0].add_run('HGVSg')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[2].paragraphs[0].add_run('HGVSc')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[3].paragraphs[0].add_run('HGVSp')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[4].paragraphs[0].add_run('Zygosity')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[5].paragraphs[0].add_run('Phenotype Contribution')
        run.bold = True
        run.font.size = Pt(9)
        run = heading_cells[6].paragraphs[0].add_run('Class')
        run.bold = True
        run.font.size = Pt(9)

    for proband_variant in proband_variants:
        cells = table.add_row().cells
        transcript = proband_variant.get_transcript()
        transcript_variant = proband_variant.get_transcript_variant()
        if transcript is None or transcript_variant is None:
            raise ValueError(f"Please select transcripts for all variants before exporting\n")
        rdr = proband_variant.create_rare_disease_report()
        run = cells[0].paragraphs[0].add_run(str(transcript.gene))
        run.font.size = Pt(7)
        run = cells[1].paragraphs[0].add_run(str(transcript_variant.hgvs_g))
        run.font.size = Pt(7)
        run = cells[2].paragraphs[0].add_run(str(transcript_variant.hgvs_c))
        run.font.size = Pt(7)
        run = cells[3].paragraphs[0].add_run(str(transcript_variant.hgvs_p))
        run.font.size = Pt(7)
        run = cells[4].paragraphs[0].add_run(str(proband_variant.zygosity))
        run.font.size = Pt(7)
        run = cells[5].paragraphs[0].add_run(str(rdr.get_contribution_to_phenotype_display()))
        run.font.size = Pt(7)
        run = cells[6].paragraphs[0].add_run(str(rdr.classification))
        run.font.size = Pt(7)

    mdt_linkage_list = MDTReport.objects.filter(interpretation_report=report).values('MDT')
    mdt = MDT.objects.filter(id__in=mdt_linkage_list).order_by('-date_of_mdt').first()

    paragraph = document.add_paragraph()

    paragraph.add_run('MDT Date: ').bold = True
    paragraph.add_run('{}\n'.format(mdt.date_of_mdt.date()))
    paragraph.add_run('MDT Attendees: ').bold = True
    clinicians = Clinician.objects.filter(mdt=mdt.id).values_list('name', flat=True)
    clinical_scientists = ClinicalScientist.objects.filter(mdt=mdt.id).values_list('name', flat=True)
    other_staff = OtherStaff.objects.filter(mdt=mdt.id).values_list('name', flat=True)

    attendees = list(clinicians) + list(clinical_scientists) + list(other_staff)
    paragraph.add_run('{}\n\n'.format(', '.join(attendees)))
    paragraph.add_run()
    run = paragraph.add_run('Discussion:\n')
    run.font.size = Pt(13)
    run.underline = True
    run.bold = True
    paragraph.add_run('{}\n\n'.format(report.ir_family.participant_family.proband.discussion.rstrip()))
    run = paragraph.add_run('Action:\n')
    run.font.size = Pt(13)
    run.underline = True
    run.bold = True
    paragraph.add_run('{}\n'.format(report.ir_family.participant_family.proband.action.rstrip()))
    return document, mdt


def write_gtab_template(report):
    '''
    Given a Cancer report, write a report template for GTAB
    :param report: GELInterpretation instance
    :return: docx document to be exported
    '''
    proband_variants = list(ProbandVariant.objects.filter(interpretation_report=report))

    document = Document()
    sections = document.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    document.add_picture(os.path.join(settings.STATIC_DIR, 'nhs_image.png'), height=Inches(0.63), width=Inches(2.39))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    heading_cells = table.rows[0].cells
    run = heading_cells[0].paragraphs[0].add_run('GTAB SUMMARY SHEET')
    run.bold = True
    heading_cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.3)
    heading_cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.3)

    table = document.add_table(rows=2, cols=1, style='Table Grid')
    table.rows[0].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'FOR RESEARCH PURPOSES ONLY- THESE RESULTS HAVE NOT BEEN VALIDATED.\n')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'UNVALIDATED FINDINGS MUST NOT BE ACTED UPON.\n')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        'PLEASE CONTACT THE LABORATORY IF VALIDATION TESTING IS REQUIRED')
    run.font.color.rgb = RGBColor(255, 0, 0)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.3)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.3)
    run = table.rows[1].cells[0].paragraphs[0].add_run('Specialist Integrated Haematological Malignancy '
                                                       'Diagnostic Service\n')
    run.font.size = Pt(6)
    run = table.rows[1].cells[0].paragraphs[0].add_run('Acquired Genomics (SIHMDS-AG), Camelia Botnar Laboratories, '
                                                       'Great Ormond Street Hospital NHS Trust,\n')
    run.font.size = Pt(6)
    run = table.rows[1].cells[0].paragraphs[0].add_run('London, WC1N 3JH. Tel: 020 7405 9200 Ex: 5755')
    run.font.size = Pt(6)
    table.rows[1].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[1].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    table.rows[1].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('GENOMICS ENGLAND PARTICIPANT INFORMATION')
    run.bold = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    table = document.add_table(rows=4, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(f'Patient Name:\t\t'
                                                       f'{report.ir_family.participant_family.proband.forename} '
                                                       f'{report.ir_family.participant_family.proband.surname}')
    run = table.rows[0].cells[1].paragraphs[0].add_run(f'GEL Participant ID:\t'
                                                       f'{report.ir_family.participant_family.proband.gel_id}')
    run = table.rows[1].cells[0].paragraphs[0].add_run(f'Gender:\t\t'
                                                       f'{report.ir_family.participant_family.proband.sex}')
    run = table.rows[1].cells[1].paragraphs[0].add_run(f'CIP ID:\t\t\t'
                                                       f'ILMN-{report.ir_family.ir_family_id}')
    run = table.rows[2].cells[0].paragraphs[0].add_run(f'Date of Birth:\t\t'
                                                       f'{report.ir_family.participant_family.proband.date_of_birth.date()}')
    run = table.rows[2].cells[1].paragraphs[0].add_run(f'NHS number:\t\t'
                                                       f'{report.ir_family.participant_family.proband.nhs_number}')
    run = table.rows[3].cells[0].paragraphs[0].add_run(f'Referring Clinician:\t'
                                                       f'{report.ir_family.participant_family.clinician.name}')
    run = table.rows[3].cells[1].paragraphs[0].add_run(f'Referring Hospital:\t'
                                                       f'{report.ir_family.participant_family.proband.gmc}')

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('GENOMICS ENGLAND REPORT DETAILS')
    run.bold = True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    table = document.add_table(rows=4, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(f'Disease Type:\t\t'
                                                       f'{report.ir_family.participant_family.proband.disease_group}')
    run = table.rows[0].cells[1].paragraphs[0].add_run(f'Tumour Content:\t\t')
    run = table.rows[1].cells[0].paragraphs[0].add_run(f'Disease Subtype:\t'
                                                       f'{report.ir_family.participant_family.proband.disease_subtype}')
    run = table.rows[1].cells[1].paragraphs[0].add_run(f'Version Number:\t\t')
    run = table.rows[2].cells[0].paragraphs[0].add_run(f'Tumour Type:\t\t')
    run = table.rows[2].cells[1].paragraphs[0].add_run(f'Total Somatic SNVs:\t\t')
    run = table.rows[3].cells[0].paragraphs[0].add_run(f'Tumour sample cross-contamination: Pass ')
    run = table.rows[3].cells[1].paragraphs[0].add_run(f'Library Prep:\t\t')
    table = document.add_table(rows=3, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('ADDITIONAL RECRUITMENT INFORMATION')
    run.bold=True
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)
    run = table.rows[1].cells[0].paragraphs[0].add_run('Multiple samples (Yes / No):   ')
    run = table.rows[2].cells[0].paragraphs[0].add_run('GENOMICS TUMOUR ADVISORY BOARD (GTAB) SUMMARY  ')
    table.rows[2].cells[0].paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.bold = True
    table.rows[2].cells[0].paragraphs[0].paragraph_format.space_before = Cm(0.2)
    table.rows[2].cells[0].paragraphs[0].paragraph_format.space_after = Cm(0.2)

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('CLINICAL UPDATE\t\t\t\t\t\t\tGTAB date: dd/mm/yyyy\n')
    run = table.rows[0].cells[0].paragraphs[0].add_run('Include any SOC testing already performed, dates of '
                                                       'treatment, clinical trials etc.\n\n\n\n\n\n\n\n\n\n')
    run.italic = True

    table = document.add_table(rows=1, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('Oncologist:')
    run = table.rows[0].cells[1].paragraphs[0].add_run('Clinician in lieu of referrer: ')

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "FOR RESEARCH PURPOSES ONLY- THESE RESULTS HAVE NOT BEEN VALIDATED\n")
    run.font.color.rgb = RGBColor(255, 0, 0)

    run = table.rows[0].cells[0].paragraphs[0].add_run("SOMATIC VARIANTS\n\n")
    run.underline=True
    run = table.rows[0].cells[0].paragraphs[0].add_run("Only variants with specific consequences "
                                                       "(transcript ablation, splice acceptor variant, splice donor "
                                                       "variant, stop gained, frameshift variant, stop lost, start lost,"
                                                       " transcript amplification, inframe insertion, inframe deletion, "
                                                       "inframe variant, missense variant, splice region variant) in "
                                                       "canonical transcripts are reported.Complex indels and "
                                                       "frameshift variants are only annotated at the CDS level owing "
                                                       "to problems accurately annotating the protein change with the "
                                                       "current pipeline.Small variants are classified as SNVs and "
                                                       "indels < 50bp.Classification for gene mode of action (oncogene, "
                                                       "tumour suppressor or both) was extracted from the manually "
                                                       "curated Cancer Gene Census list (Wellcome Trust Sanger "
                                                       "Institute).Reported variants are classified into Domains 1-3, "
                                                       "of which Domains 1-2 are reviewed by a clinical scientist and "
                                                       "discussed at GTAB.\n\n")
    run.font.size = Pt(5)
    run = table.rows[0].cells[0].paragraphs[0].add_run("Domain 1 (variants reported as having therapeutic, "
                                                       "prognostic or trial associations by GenomOncology "
                                                       "Knowledge Management System):\n")
    run.underline = True
    run.font.size = Pt(9)
    run = table.rows[0].cells[0].paragraphs[0].add_run("Note: Please see individual "
                                                       "variant for assessment of pathogenicity\n\n") #\t1)\n")
    count = 1
    for proband_variant in proband_variants:
        if proband_variant.max_tier == 1 and proband_variant.somatic is True:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n\n")
            count += 1
    table.rows[0].cells[0].paragraphs[0].add_run("\n")
    run = table.rows[0].cells[0].paragraphs[0].add_run("Domain 2 "
                                                       "(variants in genes within the Cancer Gene Census- Wellcome "
                                                       "Trust Sanger Institute):\n")
    count = 1
    for proband_variant in proband_variants:
        if proband_variant.max_tier == 2 and proband_variant.somatic is True:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n\n")
            count += 1
    run.underline=True
    run.font.size = Pt(9)
    run = table.rows[0].cells[0].paragraphs[0].add_run("Note: Please see individual "
                                                       "variant for assessment of pathogenicity\n\n")

    table = document.add_table(rows=1, cols=1, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run(
        "FOR RESEARCH PURPOSES ONLY- THESE RESULTS HAVE NOT BEEN VALIDATED\n")
    run.font.color.rgb = RGBColor(255, 0, 0)
    run = table.rows[0].cells[0].paragraphs[0].add_run("ADDITIONAL FINDINGS\n\n")
    run.underline=True
    run = table.rows[0].cells[0].paragraphs[0].add_run("Cancer Pertinent Germline Susceptibility\n\n")
    run.underline=True
    count = 1
    for proband_variant in proband_variants:
        if proband_variant.somatic is False:
            transcript = proband_variant.get_transcript()
            transcript_variant = proband_variant.get_transcript_variant()
            if transcript_variant.hgvs_c:
                hgvs_c = transcript_variant.hgvs_c.split(':')
                if len(hgvs_c) > 1:
                    hgvs_c = hgvs_c[1]
                else:
                    hgvs_c = hgvs_c[1]
            else:
                hgvs_c = None
            if transcript_variant.hgvs_p:
                hgvs_p = transcript_variant.hgvs_p.split(':')
                if len(hgvs_p) > 1:
                    hgvs_p = hgvs_p[1]
                else:
                    hgvs_p = hgvs_p[1]
            else:
                hgvs_p = None
            table.rows[0].cells[0].paragraphs[0].add_run(f"{count}) {transcript.gene} {hgvs_c} {hgvs_p} VAF: XX\n"
                                                         f"Transcript: {transcript.name}\n\n")
            count += 1
    run.italic=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Tier 1 includes variants deemed to be pathogenic or likely "
                                                        "pathogenic in cancer susceptibility genes relevant to "
                                                        "tumour type.\nTier 3 contains all rare variants arising across"
                                                        " a large set of cancer susceptibility genes. Review of Tier "
                                                        "3 germline variants is not required routinely but should be "
                                                        "considered in cases in which there is a high index of "
                                                        "suspicion of a germline determinant of cancer in the patient "
                                                        "and/or family. \n\n")
    run.font.size=Pt(5)
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Circos Plot (genome-wide visualisation of "
                                                        "somatic variants and sequencing depth)\n\n")
    run.underline=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Copy and paste image of circus plot\n\n")
    run.italic=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "This plot illustrates the distribution of "
                                                        "somatic variants across the genome with each concentric "
                                                        "circle (track) representing a different class of variant.\n"
                                                        "Chromosomes are arranged sequentially around the "
                                                        "circumference "
                                                        "as indicated. The information presented in each track is as "
                                                        "follows:\nTrack 1 (innermost track): chromosomes\nTrack 2 "
                                                        "(in red): number of somatic SNVs in 2Mb window; scale from"
                                                        " 0 to 100\nTrack 3 (in green): number of somatic indels in "
                                                        "2Mb window; scale from 0 to 35\nTrack 4: ratio of normalised "
                                                        "depth of coverage for tumour vs normal in log2 scale smoothed "
                                                        "over 100 kb windows. Diploid regions have value of 0. Scale is"
                                                        " between -2 and 2. Regions with coverage below 15x in "
                                                        "germline are not shown. CNV losses are indicated in red, "
                                                        "CNV gains are "
                                                        "indicated in green, copy-neutral LOH regions are indicated in "
                                                        "yellow.\nTrack 5 (outermost track, in blue): absolute depth "
                                                        "of coverage in tumour sample\nStructural variants (SVs) are "
                                                        "indicated by arcs inside the plot; translocations are "
                                                        "indicated in green, inversions are indicated in purple. SVs "
                                                        "shorter than 100 kb and insertions are not plotted.\n\n\n")
    run.font.size = Pt(4)
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Structural Variants\n")
    run.underline=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Translocations involving 2 or more "
                                                        "named genes:\n\n")
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Copy and paste structural variants table\n\n")
    run.italic=True

    run = table.rows[0].cells[0].paragraphs[0].add_run( "Searched for copy number variants:\n\n")
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Copy and paste structural "
                                                        "variants table\n\n")
    run.italic=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Somatic CNVs and SVs variants have not "
                                                        "been assigned to domains (see ‘Somatic Variants’ ‘Domain 1’ "
                                                        "and ‘Domain 2’ above) whilst the performance (recall and "
                                                        "precision) of the calling algorithm for CNVs and SVs is under"
                                                        " evaluation.\nOnly SVs overlapping breakends with introns or "
                                                        "exons are listed in the table above. Each row corresponds to "
                                                        "one structural variant. Types of structural variants called "
                                                        "by Canvas: GAIN(COPY NUMBER) = CNV gain, LOSS(COPY NUMBER) ="
                                                        " CNV loss, LOH(COPY NUMBER) = loss of heterozygosity. Types "
                                                        "of structural variants called by Manta: BND = translocation, "
                                                        "DEL = deletion, DUP = duplication, INV = inversion, "
                                                        "INS = insertion. Coordinate for the second breakend in "
                                                        "translocation event captures replacement string, position and "
                                                        "direction according to variant call format specification v4.3\n\n")
    run.font.size = Pt(5)
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Mutation Burden\n")
    run.underline=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Copy and paste mutation burden plot\n\n")
    run.italic=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Total "
                                                        "number of somatic non-synonymous small variants per megabase "
                                                        "(coding region): XXX\n")
    run = table.rows[0].cells[0].paragraphs[0].add_run( "The vertical axis (log scaled) shows "
                                                        "the number of somatic non-synonymous small variants per "
                                                        "megabase of coding sequences. The dashed horizontal red line "
                                                        "represents the total somatic mutation burden in this patients "
                                                        "genome. Each sample in the 100,000 Genomes Cancer dataset is "
                                                        "represented by a dot on the plot; different cancer types are "
                                                        "ordered on the horizontal axis based on their median numbers "
                                                        "of somatic mutations (short horizontal red line).\n\n\n")
    run.font.size = Pt(5)
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Mutational Signature\n")
    run.underline=True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Copy and paste mutational signatures\n\n")
    run.italic = True
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Further "
                                                        "details of the 30 different mutational signatures used for "
                                                        "this analysis, their prevalence in different tumour types "
                                                        "and proposed aetiology can be found at the Sanger Institute "
                                                        "Website- https://cancer.sanger.ac.uk/cosmic/signatures\n\n")
    run.font.size = Pt(5)
    run = table.rows[0].cells[0].paragraphs[0].add_run( "Pharmacogenomics\n\n")
    run.underline=True

    table = document.add_table(rows=2, cols=2, style='Table Grid')
    run = table.rows[0].cells[0].paragraphs[0].add_run('Completed by:')
    run = table.rows[0].cells[1].paragraphs[0].add_run('Date: ')
    run = table.rows[1].cells[0].paragraphs[0].add_run('Checked by:')
    run = table.rows[1].cells[1].paragraphs[0].add_run('Date: ')

    return document
